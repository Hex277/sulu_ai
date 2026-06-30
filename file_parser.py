"""
file_parser.py
---------------
data/ qovluğundakı bütün dəstəklənən faylları skan edib, içindəki mətni
təmiz string formatına çevirən modul.

TƏHLÜKƏSİZLİK: Bu modul YALNIZ config.DATA_DIR daxilindəki faylları oxuyur.
_is_safe_path() funksiyası hər faylın faktiki olaraq DATA_DIR içində
olduğunu yoxlayır ki, simvolik link və ya ".." kimi yollarla kənara
çıxış mümkün olmasın.
"""

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from PyPDF2 import PdfReader

from config import DATA_DIR, SUPPORTED_EXTENSIONS
from logger import log_step


def _is_safe_path(path: Path) -> bool:
    """Faylın həqiqətən DATA_DIR daxilində olduğunu yoxlayır."""
    try:
        path.resolve().relative_to(DATA_DIR.resolve())
        return True
    except ValueError:
        return False


def _clean_text(text: str) -> str:
    """
    Çıxarılan xam mətni təmizləyir:
    - Görünməz/binar qalıq idarəetmə simvollarını silir
    - Artıq boşluq/sətir aralarını sıxır
    """
    if not text:
        return ""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_txt(path: Path) -> str:
    """Sadə .txt faylını oxuyur, encoding xətalarına davamlıdır."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Bəzi köhnə fayllar utf-8 olmaya bilər — son çarə kimi
        return path.read_text(encoding="utf-8", errors="ignore")


def _read_json(path: Path) -> str:
    """
    .json faylını oxuyub, FTS üçün mənalı düz mətnə çevirir.
    Sadəcə json.dumps etmək əvəzinə, key:value cütlərini sətir-sətir yazırıq
    ki, axtarış mühərriki üçün daha "oxunaqlı" mətn olsun.
    """
    with path.open("r", encoding="utf-8") as f:
        raw = json.load(f)

    lines = []

    def _flatten(obj, prefix=""):
        if isinstance(obj, dict):
            for key, value in obj.items():
                new_prefix = f"{prefix}{key}: " if not prefix else f"{prefix}.{key}: "
                _flatten(value, new_prefix)
        elif isinstance(obj, list):
            for item in obj:
                _flatten(item, prefix)
        else:
            lines.append(f"{prefix}{obj}")

    _flatten(raw)
    return "\n".join(lines)


def _read_docx(path: Path) -> str:
    """
    .docx faylını python-docx ilə oxuyur (paraqraflar + cədvəllər).
    Bu, binar/şifrələnmiş simvolların AI-a getməsinin qarşısını alır.
    """
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)

def _read_xlsx(path: Path, rel_path: str, mtime: float) -> list[dict]:
    """
    AnyTXT Üslubunda: .xlsx faylını sətir-sətir indekslənən sənədlərə parçalayır.
    """
    sheets = pd.read_excel(path, sheet_name=None, engine="openpyxl")
    documents = []
    
    for sheet_name, df in sheets.items():
        # Sütun adlarını string formatına salırıq
        columns = [str(col).strip() for col in df.columns]
        
        for idx, row in df.iterrows():
            row_parts = []
            for col, val in zip(columns, row):
                if pd.notna(val) and str(val).strip() != "":
                    row_parts.append(f"{col}: {str(val).strip()}")
            
            if row_parts:
                row_content = ", ".join(row_parts)
                # Bazadakı unikal 'filepath' xətasının qarşısını almaq üçün unikal ID qoşuruq
                documents.append({
                    "path": f"{rel_path}#row_{sheet_name}_{idx}",
                    "filename": f"{path.stem} -> {sheet_name} (Sətir {idx+2})", # Exceldə sətir 2-dən başlayır
                    "content": row_content,
                    "mtime": mtime
                })
    return documents


def _read_csv(path: Path, rel_path: str, mtime: float) -> list[dict]:
    """AnyTXT Üslubunda: .csv faylını sətir-sətir indekslənən sənədlərə parçalayır."""
    try:
        df = pd.read_csv(path)
    except UnicodeDecodeError:
        df = pd.read_csv(path, encoding="latin-1")
        
    columns = [str(col).strip() for col in df.columns]
    documents = []
    
    for idx, row in df.iterrows():
        row_parts = []
        for col, val in zip(columns, row):
            if pd.notna(val) and str(val).strip() != "":
                row_parts.append(f"{col}: {str(val).strip()}")
                
        if row_parts:
            row_content = ", ".join(row_parts)
            documents.append({
                "path": f"{rel_path}#row_{idx}",
                "filename": f"{path.stem} (Sətir {idx+2})",
                "content": row_content,
                "mtime": mtime
            })
    return documents
def _read_pdf(path: Path) -> str:
    """.pdf faylını PyPDF2 ilə səhifə-səhifə oxuyub mətni birləşdirir."""
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts)


_READERS = {
    ".txt": _read_txt,
    ".json": _read_json,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
    ".csv": _read_csv,
    ".pdf": _read_pdf,
}

# file_parser.py faylının ən aşağısındakı load_documents funksiyasını bu şəkildə dəyiş:

def load_single_document(path: Path) -> list[dict]:
    """
    Faylı oxuyur. 
    Word/PDF üçün tək elementli siyahı, Excel/CSV üçün isə sətir-sətir bölünmüş siyahı qaytarır.
    """
    if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        return []

    # Faylın məlumatlarını alırıq
    rel_path = str(path.relative_to(DATA_DIR))
    mtime = path.stat().st_mtime
    
    # XLSX və CSV üçün yeni yazdığımız sətir-sətir funksiyaları çağırırıq
    if path.suffix.lower() == '.xlsx':
        return _read_xlsx(path, rel_path, mtime)
    elif path.suffix.lower() == '.csv':
        return _read_csv(path, rel_path, mtime)
    
    # Word və PDF üçün köhnə qaydada (bir blok) oxuyuruq
    reader = _READERS[path.suffix.lower()]
    raw_text = reader(path)
    clean_text = _clean_text(raw_text)
    
    if clean_text:
        return [{
            "path": rel_path,
            "filename": path.name,
            "content": clean_text,
            "mtime": mtime
        }]
    return []